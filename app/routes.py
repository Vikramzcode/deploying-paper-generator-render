import os
import json
import re
import uuid
import io
from datetime import datetime
from flask import Blueprint, request, jsonify, render_template, current_app, send_file, render_template_string, make_response
from sqlalchemy.sql.expression import func
from docx import Document
from weasyprint import HTML
import pathlib
# Keep reportlab imports ONLY for the answer key generation
from reportlab.lib.pagesizes import A4
from reportlab.platypus import Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

import google.generativeai as genai
from .models import Question, Paper, PaperQuestion, Visitor
from . import db

main = Blueprint("main", __name__)

def get_or_create_visitor():
    """Get existing visitor from cookie or create new visitor"""
    visitor_id = request.cookies.get('visitor_id')
    
    if visitor_id:
        # Check if visitor exists in database
        visitor = Visitor.query.filter_by(visitor_id=visitor_id).first()
        if visitor:
            # Update last visit and increment visit count
            visitor.visit_count += 1
            db.session.commit()
            return visitor
    
    # Create new visitor
    visitor_id = str(uuid.uuid4())
    visitor = Visitor()
    visitor.visitor_id = visitor_id
    db.session.add(visitor)
    db.session.commit()
    
    return visitor

@main.route("/")
def index():
    response = make_response(render_template("index.html"))
    
    # Check if visitor cookie exists
    visitor_id = request.cookies.get('visitor_id')
    if not visitor_id:
        # Create new visitor and set cookie
        visitor = get_or_create_visitor()
        response.set_cookie('visitor_id', visitor.visitor_id, max_age=365*24*60*60)  # 1 year
    
    return response

# Helper function to normalize question types
def _normalize_qtype(label):
    label = (label or "").strip()
    if label in ["MCQ", "Multiple Choice"]: return "MCQ"
    if label in ["Fill in the Blanks", "Fill"]: return "Fill in the Blanks"
    if label in ["Short Answer", "Short"]: return "Short Answer"
    if label in ["Long Answer", "Long"]: return "Long Answer"
    if label in ["Matching", "Match", "Match the Following"]: return "Matching"
    if label in ["Case Study", "Case"]: return "Case Study"
    return label

@main.route("/api/generate", methods=["POST"])
def generate_paper():
    data = request.get_json()
    current_app.logger.info(f"Incoming /api/generate payload: {data}")

    subject = data.get("subject")
    class_ = data.get("class")
    board = data.get("schoolBoard")
    school = data.get("schoolName")
    qdist = data.get("questionDistribution", {})
    ddist = data.get("difficultyDistribution", {})
    exam_name = data.get("examName")
    paper_language = data.get("paperLanguage", "english")
    topic = data.get("topic", "")  # Get topic if provided
    chapters = data.get("chapters", [])  # Get chapters if provided
    questions = []

    # Get or create visitor
    visitor = get_or_create_visitor()
    
    # More robust topic detection
    topic_present = bool(topic and topic.strip())
    current_app.logger.info(f"Topic received: '{topic}'")
    current_app.logger.info(f"Topic stripped: '{topic.strip()}'")
    current_app.logger.info(f"Topic bool check: {topic_present}")
    current_app.logger.info(f"Using topic-based generation: {topic_present}")
    
    # Log all the data to debug
    current_app.logger.info(f"All data received: {data}")
    
    # Test if API key is working
    try:
        current_app.logger.info("Testing Google Generative AI API key")
        model = genai.GenerativeModel('models/gemini-flash-latest')
        test_response = model.generate_content("Say 'Hello, World!' in one word.")
        current_app.logger.info(f"API key test response: {test_response.text.strip()}")
    except Exception as e:
        current_app.logger.error(f"API key test failed: {e}")
    
    qdist_str_parts = []
    for qtype, info in qdist.items():
        count = info.get('count', 0)
        if count > 0:
            qdist_str_parts.append(f"- {count} {qtype} question(s)")
    qdist_prompt_str = "\n".join(qdist_str_parts)
    
    language_instruction = ""
    if paper_language == "hindi" and subject.lower() != "english":
        language_instruction = "You MUST generate the entire question paper, including questions, options, answers, and explanations, strictly in Hindi using Unicode characters."
    elif paper_language == "english" or subject.lower() == "english":
        language_instruction = "You MUST generate the entire question paper, including questions, options, answers, and explanations, strictly in English."

    try:
        # Create different prompts based on generation mode
        # More robust topic detection
        topic_present = bool(topic and topic.strip())
        current_app.logger.info(f"Using topic-based generation: {topic_present}")
        
        if topic_present:
            # Topic-based generation - improved prompt
            current_app.logger.info("Generating topic-based prompt")
            prompt = f"""
You are an experienced {board} school teacher creating a question paper. 
Your task is to create a comprehensive question paper focused EXCLUSIVELY on the topic: "{topic}".

School: {school}
Class: {class_}
Subject: {subject}

{language_instruction}

CRITICAL INSTRUCTIONS:
1. Generate EXACTLY the number and types of questions specified below
2. ALL questions MUST be directly related to the topic "{topic}"
3. Do NOT include any questions unrelated to this topic
4. Follow the difficulty distribution as closely as possible
5. ALL content MUST be in {paper_language.capitalize()} language

Question Distribution Requirements:
{qdist_prompt_str}

Difficulty Distribution Target:
{ddist}

OUTPUT FORMAT REQUIREMENTS:
- Return ONLY a valid JSON array of question objects
- Do NOT include any other text, explanations, or markdown formatting
- Each question object MUST have these exact keys:
  - "type" (e.g., "MCQ", "Short Answer", "Long Answer", etc.)
  - "question" (the question text)
  - "options" (ONLY for "MCQ" type: exactly 4 options as a JSON array of strings)
  - "marks" (integer)
  - "difficulty" ("Easy", "Medium", or "Hard")
  - "answer" (correct answer - for MCQ provide the letter like "A" or "B")
  - "explanation" (brief explanation)

Example MCQ format:
{{
  "type": "MCQ",
  "question": "What is the capital of France?",
  "options": ["London", "Berlin", "Paris", "Madrid"],
  "marks": 1,
  "difficulty": "Easy",
  "answer": "C",
  "explanation": "Paris is the capital of France."
}}

Begin generating the question paper now:
"""
        else:
            # Class and chapters-based generation
            current_app.logger.info("Generating class/chapter-based prompt")
            chapters_str = ""
            if chapters:
                chapters_str = f"Focus on these chapters: {', '.join(chapters)}."
            
            prompt = f"""
You are an experienced {board} school teacher creating a question paper. 
Your task is to create a comprehensive question paper for Class {class_}, Subject: {subject}.

School: {school}
{chapters_str}

{language_instruction}

CRITICAL INSTRUCTIONS:
1. Generate EXACTLY the number and types of questions specified below
2. ALL questions MUST be appropriate for Class {class_} {subject}
3. Follow the difficulty distribution as closely as possible
4. ALL content MUST be in {paper_language.capitalize()} language

Question Distribution Requirements:
{qdist_prompt_str}

Difficulty Distribution Target:
{ddist}

OUTPUT FORMAT REQUIREMENTS:
- Return ONLY a valid JSON array of question objects
- Do NOT include any other text, explanations, or markdown formatting
- Each question object MUST have these exact keys:
  - "type" (e.g., "MCQ", "Short Answer", "Long Answer", etc.)
  - "question" (the question text)
  - "options" (ONLY for "MCQ" type: exactly 4 options as a JSON array of strings)
  - "marks" (integer)
  - "difficulty" ("Easy", "Medium", or "Hard")
  - "answer" (correct answer - for MCQ provide the letter like "A" or "B")
  - "explanation" (brief explanation)

Example MCQ format:
{{
  "type": "MCQ",
  "question": "What is the capital of France?",
  "options": ["London", "Berlin", "Paris", "Madrid"],
  "marks": 1,
  "difficulty": "Easy",
  "answer": "C",
  "explanation": "Paris is the capital of France."
}}

Begin generating the question paper now:
"""
            
        # Log the prompt to see what's being sent to the AI
        current_app.logger.info(f"Sending prompt to AI: {prompt}")
        
        # Use the already configured genai from __init__.py
        # Using getattr to avoid linter issues
        try:
            current_app.logger.info("Attempting to create GenerativeModel with 'models/gemini-flash-latest'")
            model = genai.GenerativeModel("models/gemini-flash-latest")  # Using the latest flash model
            current_app.logger.info("Model created successfully, sending prompt")
            response = model.generate_content(prompt)
            raw_text = response.text.strip()
            current_app.logger.info("Response received successfully")
        except Exception as model_error:
            current_app.logger.error(f"AI model error with models/gemini-flash-latest: {model_error}")
            # Try a fallback model
            try:
                current_app.logger.info("Attempting fallback with 'models/gemini-pro-latest'")
                model = genai.GenerativeModel("models/gemini-pro-latest")  # Fallback to latest pro model
                response = model.generate_content(prompt)
                raw_text = response.text.strip()
                current_app.logger.info("Fallback response received successfully")
            except Exception as fallback_error:
                current_app.logger.error(f"AI model error with fallback models/gemini-pro-latest: {fallback_error}")
                # Try another fallback model
                try:
                    current_app.logger.info("Attempting fallback with 'models/gemini-2.0-flash'")
                    model = genai.GenerativeModel("models/gemini-2.0-flash")  # Another fallback option
                    response = model.generate_content(prompt)
                    raw_text = response.text.strip()
                    current_app.logger.info("Second fallback response received successfully")
                except Exception as second_fallback_error:
                    current_app.logger.error(f"AI model error with second fallback models/gemini-2.0-flash: {second_fallback_error}")
                    raise second_fallback_error
        
        # Log the AI response
        current_app.logger.info(f"AI response: {raw_text}")

        try:
            questions = json.loads(raw_text)
        except json.JSONDecodeError:
            json_match = re.search(r"```json\s*(.*?)```", raw_text, re.DOTALL)
            if json_match:
                questions = json.loads(json_match.group(1))
            else:
                raise ValueError("AI returned invalid JSON.")

        for q in questions:
            q["question_type"] = _normalize_qtype(q.get("type"))
        
        # Track used question IDs to prevent duplicates
        used_question_hashes = set()
        
        balanced_questions = []
        for qtype_frontend, info in qdist.items():
            count_needed = int(info['count'])
            normalized_type = _normalize_qtype(qtype_frontend)
            # Filter questions of this type
            type_questions = [q for q in questions if q.get("question_type") == normalized_type]
            
            # Add questions to balanced list, avoiding duplicates
            added_count = 0
            for q in type_questions:
                # Check if we've already added enough questions of this type
                if added_count >= count_needed:
                    break
                    
                # Create a simple hash of the question text to detect duplicates
                question_hash = hash(q.get("question", "").strip().lower())
                if question_hash not in used_question_hashes:
                    balanced_questions.append(q)
                    used_question_hashes.add(question_hash)
                    added_count += 1
        
        questions = balanced_questions
        
        # Track used question texts to prevent duplicates in processed questions
        used_question_texts = set()
        processed_questions = []
        for q in questions:
            q_type = q.get("question_type")
            q_text = q.get("question", "")
            
            # Skip duplicate questions
            q_text_normalized = q_text.strip().lower()
            if q_text_normalized in used_question_texts:
                continue
            used_question_texts.add(q_text_normalized)
            
            options = q.get("options", [])
            
            if q_type == "MCQ":
                if not isinstance(options, list): options = []
                while len(options) < 4: options.append(f"Option {len(options)+1}")
                options = options[:4]
            
            # Fix for Question model instantiation - using keyword arguments correctly
            new_q = Question()
            new_q.school_name = school
            new_q.board = board
            new_q.class_ = class_
            new_q.subject = subject
            new_q.question_type = q_type
            new_q.difficulty = q.get("difficulty")
            new_q.marks = q.get("marks")
            new_q.question_text = q_text
            new_q.answer = q.get("answer", "Not provided")
            new_q.source = "AI"
            new_q.explanation = q.get("explanation", "")
            new_q.options = options if q_type == "MCQ" else None
            new_q.language = paper_language
            
            db.session.add(new_q)
            db.session.flush()

            q['id'] = new_q.id
            q['source'] = "AI"
            q['options'] = options if q_type == "MCQ" else None
            q['question_text'] = q_text
            processed_questions.append(q)

        questions = processed_questions
        db.session.commit()

    except Exception as e:
        current_app.logger.error(f"AI generation failed: {e}. Falling back completely to DB.")
        db.session.rollback()
        questions = []

    # --- Fallback: fill missing questions from DB ---
    total_needed = sum(int(info['count']) for info in qdist.values())
    # Create a set of existing question texts to avoid duplicates
    existing_question_texts = {q.get("question_text", "").strip().lower() for q in questions}
    
    if len(questions) < total_needed:
        for qtype_frontend, info in qdist.items():
            count_needed = int(info['count'])
            marks = int(info['marks'])
            
            # --- FIX #4: USE NORMALIZED TYPE FOR COUNTING AND QUERYING ---
            normalized_type = _normalize_qtype(qtype_frontend)
            
            # Count how many questions of this type we already have
            picked_count = sum(1 for q in questions if q.get("question_type") == normalized_type)
            missing_for_type = max(0, count_needed - picked_count)

            if missing_for_type > 0:
                # Create the base query with exact criteria ONLY
                query = Question.query.filter_by(
                    question_type=normalized_type, 
                    marks=marks, 
                    language=paper_language
                )
                
                # Add topic-based filtering if topic is provided
                if topic_present:
                    # For topic-based generation, we'll search for questions that might contain the topic
                    query = query.filter(Question.question_text.contains(topic))
                else:
                    # For class/chapter-based generation, use the existing filters
                    if subject:
                        query = query.filter_by(subject=subject)
                    if class_:
                        query = query.filter_by(class_=class_)
                
                # Get more questions than needed to filter duplicates
                db_questions = (
                    query
                    .order_by(func.rand())
                    .limit(missing_for_type * 2)  # Get extra questions to filter duplicates
                    .all()
                )
                
                # Add questions while avoiding duplicates
                added_count = 0
                for q in db_questions:
                    if added_count >= missing_for_type:
                        break
                        
                    question_text = q.question_text.strip().lower()
                    if question_text not in existing_question_texts:
                        q_dict = q.as_dict()
                        q_dict['source'] = "Database"
                        # Ensure language is preserved when fetching from database
                        q_dict['language'] = q.language
                        questions.append(q_dict)
                        existing_question_texts.add(question_text)
                        added_count += 1
                        
                # Log how many questions we were able to add
                current_app.logger.info(f"Added {added_count} questions of type {normalized_type} from database. Total questions now: {len(questions)}")
        
        # Final fallback - if we still don't have enough questions, generate some basic ones
        if len(questions) < total_needed:
            current_app.logger.info(f"Still need {total_needed - len(questions)} questions, generating basic fallback questions")
            missing_count = total_needed - len(questions)
            
            # Generate some basic questions as a last resort
            for i in range(missing_count):
                basic_question = {
                    "type": "Short",
                    "question_type": "Short Answer",
                    "question": f"Explain the basic concepts related to {subject} for Class {class_}.",
                    "marks": 3,
                    "difficulty": "Medium",
                    "answer": "Answer would depend on the specific topic.",
                    "explanation": "This is a fallback question generated due to limited database content.",
                    "source": "Fallback",
                    "language": paper_language,
                    "options": None
                }
                questions.append(basic_question)
                
            current_app.logger.info(f"Generated {missing_count} fallback questions. Total questions now: {len(questions)}")
    paper_id = str(uuid.uuid4())[:8]
    pdf_filename = f"paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    word_filename = f"{paper_id}.docx"
    answer_key_filename = f"answer_key_{paper_id}.pdf"

    paper_entry = Paper()
    paper_entry.paper_id = paper_id
    paper_entry.exam_name = exam_name
    paper_entry.school_name = school
    paper_entry.board = board
    paper_entry.class_ = class_
    paper_entry.subject = subject
    paper_entry.total_questions = len(questions)
    paper_entry.total_marks = sum(int(q.get("marks", 0)) for q in questions)
    paper_entry.pdf_path = f"/static/papers/{pdf_filename}"
    paper_entry.word_path = f"/static/papers/{word_filename}"
    paper_entry.answer_key_path = f"/static/papers/{answer_key_filename}"
    paper_entry.visitor_id = visitor.visitor_id  # Link paper to visitor
    
    db.session.add(paper_entry)
    db.session.commit()

    for q in questions:
        pq = PaperQuestion()
        pq.paper_id = paper_id
        pq.question_id = q['id']
        pq.question_text = q.get("question_text")
        pq.type = q.get("question_type")
        pq.difficulty = q.get("difficulty")
        pq.marks = q.get("marks")
        pq.options = q.get("options") if q.get("options") else None
        pq.answer = q.get("answer", "Not provided")
        db.session.add(pq)
    db.session.commit()

    papers_dir = os.path.join(current_app.root_path, "static", "papers")
    os.makedirs(papers_dir, exist_ok=True)
    json_path = os.path.join(papers_dir, f"{paper_id}.json")

    summary = {
        "total_questions": len(questions),
        "total_marks": sum(int(q.get("marks", 0)) for q in questions)
    }

    type_order = ["MCQ", "Multiple Choice", "Fill in the Blanks", "Fill", "Short Answer", "Short", "Long Answer", "Long", "Matching", "Match", "Match the Following", "Case Study", "Case"]
    def get_type_order(q):
        t = (q.get("type") or q.get("question_type") or "").strip()
        try: return type_order.index(t)
        except ValueError: return len(type_order)
    questions_sorted = sorted(questions, key=get_type_order)

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump({
            "paper_id": paper_id, "examName": exam_name, "schoolName": school,
            "schoolBoard": board, "class": class_, "subject": subject,
            "questions": questions_sorted, "summary": summary
        }, f, indent=2, ensure_ascii=False)


    # --- PDF GENERATION WITH WEASYPRINT (FINAL VERSION) ---
    
    # Get the absolute path to the local font file
    font_path = os.path.join(current_app.root_path, 'fonts', 'NotoSansDevanagari-Regular.ttf')
    font_url = pathlib.Path(font_path).as_uri() # Converts path to file:/// URI

    sections = {
        "Multiple Choice": [], "Fill in the Blanks": [], "Short Answer": [],
        "Long Answer": [], "Matching": [], "Case Study": []
    }
    section_titles = {
        "Multiple Choice": "Section A - Multiple Choice Questions",
        "Fill in the Blanks": "Section B - Fill in the Blanks",
        "Short Answer": "Section C - Short Answer Questions",
        "Long Answer": "Section D - Long Answer Questions",
        "Matching": "Section E - Matching Questions",
        "Case Study": "Section F - Case Study"
    }
    for q in questions_sorted:
        q_type = _normalize_qtype(q.get("question_type", ""))
        if q_type == "MCQ": q_type = "Multiple Choice"
        if q_type in sections:
            sections[q_type].append(q)

    # Define the HTML template for the PDF
    # ...existing code...

    html_template = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {
                size: A4;
                margin: 25mm 20mm 25mm 20mm;
            }
            @font-face {
                font-family: 'Noto Sans Devanagari';
                src: url('{{ font_url }}');
            }
            html, body {
                font-family: 'Noto Sans Devanagari', sans-serif;
                font-size: 12pt;
                background: #fff;
            }
            .paper-container {
                width: 100%;
                margin: 0;
                padding: 0;
            }
            .paper-header {
                text-align: center;
                margin-bottom: 18px;
                border-bottom: 2px solid #222;
                padding-bottom: 8px;
                page-break-after: avoid;
            }
            .paper-header h1 { font-size: 22pt; margin: 0; }
            .paper-header h2 { font-size: 16pt; margin: 4px 0; font-weight: normal; }
            .paper-header h3 { font-size: 13pt; margin: 4px 0; font-weight: normal; }
            .details {
                display: flex;
                justify-content: space-between;
                margin-bottom: 18px;
                font-size: 11pt;
            }
            section {
                margin-bottom: 22px;
                page-break-inside: avoid;
            }
            h4.section-title {
                font-size: 13pt;
                border-bottom: 1px solid #bbb;
                padding: 4px 0;
                margin-bottom: 12px;
                margin-top: 0;
            }
            ol.question-list {
                list-style-type: none;
                padding-left: 0;
                margin-top: 0;
            }
            li.question {
                margin-bottom: 16px;
            }
            .question-text {
                display: flex;
                justify-content: space-between;
            }
            .question-text .marks {
                font-weight: bold;
                white-space: nowrap;
                padding-left: 12px;
            }
            ol.options {
                list-style-type: lower-alpha;
                padding-left: 32px;
                margin-top: 7px;
                margin-bottom: 0;
            }
            .option {
                margin-bottom: 4px;
            }
        </style>
    </head>
    <body>
        <div class="paper-container">
            <div class="paper-header">
                <h1>{{ school }}</h1>
                <h2>{{ exam_name }}</h2>
                <h3>Class {{ class_ }} - {{ subject }}</h3>
            </div>
            <div class="details">
                <span>Date: {{ date }}</span>
            </div>

            {% set q_num = namespace(value=1) %}
            {% for sec_type, q_list in sections.items() %}
                {% if q_list %}
                <section>
                    <h4 class="section-title">{{ section_titles[sec_type] }}</h4>
                    <ol class="question-list">
                        {% for q in q_list %}
                        <li class="question">
                            <div class="question-text">
                                <span><b>Q{{ q_num.value }}.</b> {{ q.question_text }}</span>
                                <span class="marks">({{ q.marks }} marks)</span>
                            </div>
                            {% if q.question_type == 'MCQ' and q.options %}
                            <ol class="options">
                                {% for opt in q.options %}
                                <li class="option">{{ opt }}</li>
                                {% endfor %}
                            </ol>
                            {% endif %}
                        </li>
                        {% set q_num.value = q_num.value + 1 %}
                        {% endfor %}
                    </ol>
                </section>
                {% endif %}
            {% endfor %}
        </div>
    </body>
    </html>
    """
# ...existing code...

    
    def to_char_filter(n):
        return chr(97 + n)  # Convert to lowercase letters (a, b, c, d)
    
    env = current_app.jinja_env
    env.filters['to_char'] = to_char_filter
    
    rendered_html = render_template_string(
        html_template, 
        font_url=font_url, # Pass the local font path to the template
        school=school,
        exam_name=exam_name or f"{board} Board Examination",
        class_=class_,
        subject=subject,
        date=datetime.now().strftime('%d-%m-%Y'),
        sections=sections,
        section_titles=section_titles
    )
    
    pdf_path = os.path.join(papers_dir, pdf_filename)
    HTML(string=rendered_html).write_pdf(pdf_path)


    return jsonify({
        "questions": [{
            "id": q.get("id"), "question_text": q.get("question_text"),
            "marks": q.get("marks"), "difficulty": q.get("difficulty"),
            "type": q.get("type") or q.get("question_type"), "source": q.get("source", "Database")
        } for q in questions_sorted],
        "summary": summary,
        "pdf_url": f"/static/papers/{pdf_filename}",
        "word_url": f"/api/download/word/{paper_id}",
        "answer_key_url": f"/api/download/answer_key/{paper_id}"
    })


@main.route("/api/download/word/<paper_id>", methods=["GET"])
def download_word(paper_id):
    json_path = os.path.join(current_app.root_path, "static", "papers", f"{paper_id}.json")
    if not os.path.exists(json_path):
        return jsonify({"error": "Paper not found"}), 404
    with open(json_path, "r", encoding="utf-8") as f:
        paper = json.load(f)
    doc = Document()
    doc.add_heading(paper.get("examName", "Question Paper"), 0)
    doc.add_paragraph(f"School: {paper.get('schoolName')}")
    doc.add_paragraph(f"Board: {paper.get('schoolBoard')}")
    doc.add_paragraph(f"Class: {paper.get('class')}  Subject: {paper.get('subject')}")
    doc.add_heading("Questions", level=1)
    for i, q in enumerate(paper.get("questions", []), 1):
        doc.add_paragraph(f"Q{i}. {q['question_text']} ({q['marks']} marks) [{q['difficulty']}]")
        if q.get("question_type") in ["MCQ", "Multiple Choice"]:
            options = q.get("options", [])
            for idx, opt in enumerate(options, start=1):
                doc.add_paragraph(f"   ({chr(96+idx)}) {opt}", style="List Bullet")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=f"paper_{paper_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@main.route("/api/download/answer_key/<paper_id>", methods=["GET"])
def download_answer_key(paper_id):
    json_path = os.path.join(current_app.root_path, "static", "papers", f"{paper_id}.json")
    if not os.path.exists(json_path):
        return jsonify({"error": "Paper not found"}), 404

    with open(json_path, "r", encoding="utf-8") as f:
        paper = json.load(f)

    # NOTE: This still uses reportlab. It may have issues with Hindi rendering.
    pdfmetrics.registerFont(TTFont("NotoSans", os.path.join(current_app.root_path, "fonts", "NotoSansDevanagari-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("NotoSans-Bold", os.path.join(current_app.root_path, "fonts", "NotoSansDevanagari-Bold.ttf")))

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    
    styles = getSampleStyleSheet()
    styleN = styles["Normal"]
    styleN.fontName = "NotoSans"
    styleB = styles["h5"]
    styleB.fontName = "NotoSans-Bold"
    
    c.setFont("NotoSans-Bold", 16)
    c.drawCentredString(width/2, height - 50, f"Answer Key - {paper.get('examName', 'Exam')}")
    c.setFont("NotoSans", 12)
    c.drawString(50, height - 80, f"School: {paper.get('schoolName', '')}")
    c.drawString(50, height - 100, f"Class: {paper.get('class', '')} | Subject: {paper.get('subject', '')}")
    c.line(50, height - 110, width - 50, height - 110)

    y = height - 140
    for i, q in enumerate(paper.get("questions", []), 1):
        question_text = q.get("question_text", "")
        answer_text = q.get("answer", "Answer not available")
        explanation_text = q.get("explanation", "")

        p_q = Paragraph(f"<b>Q{i}:</b> {question_text}", styleN)
        w, h = p_q.wrap(width - 100, y)
        if y - h < 50:
            c.showPage()
            y = height - 50
        p_q.drawOn(c, 50, y - h)
        y -= (h + 10)

        p_a = Paragraph(f"<b>Answer:</b> {answer_text}", styleN)
        w, h = p_a.wrap(width - 100, y)
        if y - h < 50:
            c.showPage()
            y = height - 50
        p_a.drawOn(c, 70, y - h)
        y -= (h + 10)

        if explanation_text:
            p_exp = Paragraph(f"<b>Explanation:</b> {explanation_text}", styleN)
            w, h = p_exp.wrap(width - 100, y)
            if y - h < 50:
                c.showPage()
                y = height - 50
            p_exp.drawOn(c, 70, y - h)
            y -= (h + 20)
        else:
            y -= 10

    c.save()
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name=f"answer_key_{paper_id}.pdf",
        mimetype="application/pdf"
    )
